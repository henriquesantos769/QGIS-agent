from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
import geopandas as gpd
from pathlib import Path
import math


def fmt_coord(v):
    return f"{v:,.4f}".replace(",", "X").replace(".", ",").replace("X", ".")

def fmt_dist(v):
    return f"{v:,.2f}".replace(".", ",")

def azimute_dms(graus):
    g = int(graus)
    m_float = abs((graus - g) * 60)
    m = int(m_float)
    s = (m_float - m) * 60
    return f"{g}°{m:02d}'{int(s):02d}\""

def _fmt_num_br(x, casas=2):
    """Formata número float no padrão brasileiro, com vírgula."""
    if x is None:
        return ""
    return f"{x:.{casas}f}".replace(".", ",")

def _fmt_coord(x, casas=4):
    """Formata coordenada com 4 casas decimais, padrão BR."""
    return _fmt_num_br(x, casas=casas)

def _add_cabecalho_memorial(doc: Document,
                            titulo="MEMORIAL DESCRITIVO",
                            quadra=None,
                            nucleo=None,
                            municipio=None,
                            uf=None,
                            promotor=None):
    """Monta o cabeçalho padrão do memorial."""
    p = doc.add_paragraph()
    run = p.add_run(titulo)
    run.bold = True
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if quadra is not None and nucleo and municipio and uf:
        p2 = doc.add_paragraph()
        texto_quad = f"QUADRA {quadra} - NÚCLEO {nucleo} – {municipio} - {uf}"
        run2 = p2.add_run(texto_quad)
        run2.bold = True
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # linha em branco

    if promotor or nucleo or municipio or uf:
        t = doc.add_paragraph()
        if promotor:
            t.add_run("Promotor da REURB: ").bold = True
            t.add_run(str(promotor) + "\n")
        if nucleo:
            t.add_run("Núcleo: ").bold = True
            t.add_run(str(nucleo) + "\n")
        if municipio or uf:
            t.add_run("Local: ").bold = True
            loc = municipio or ""
            if uf:
                loc += f" - {uf}"
            t.add_run(loc)

def add_cabecalho_memorial_quadras(doc: Document,
                                   nucleo: str = "Teste",
                                   municipio: str = "Teste",
                                   uf: str = "Teste",
                                   promotor: str = "Instituto Cidade Legal"):
    """
    Adiciona o cabeçalho formatado exatamente como o modelo do memorial das quadras.
    """

    # -----------------------------------------
    # TÍTULO PRINCIPAL
    # -----------------------------------------
    p_title = doc.add_paragraph()
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p_title.add_run("MEMORIAL DESCRITIVO")
    run.bold = True
    run.font.size = Pt(14)

    # -----------------------------------------
    # SUBTÍTULO (NÚCLEO + MUNICÍPIO)
    # -----------------------------------------
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p_sub.add_run(f"NÚCLEO {nucleo.upper()} - {municipio.upper()} - {uf}")
    run.bold = True
    run.font.size = Pt(12)

    # Espaço
    doc.add_paragraph()

    # -----------------------------------------
    # BLOCO DE INFORMAÇÕES (ALINHADO À ESQUERDA)
    # -----------------------------------------
    p_info = doc.add_paragraph()
    p_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    run = p_info.add_run("Promotor\nda REURB:   ")
    run.bold = True
    run.font.size = Pt(11)
    p_info.add_run(f"{promotor}\n")

    run = p_info.add_run("Núcleo:       ")
    run.bold = True
    run.font.size = Pt(11)
    p_info.add_run(f"{nucleo}\n")

    run = p_info.add_run("Local:        ")
    run.bold = True
    run.font.size = Pt(11)
    p_info.add_run(f"{municipio} - {uf}\n")

    # Espaço grande antes da descrição
    doc.add_paragraph()

def add_bloco_info_quadra(doc, quadra, area_m2, perimetro_m):
    """
    Adiciona o bloco inicial da quadra:
    
    QUADRA  XX
    Área:      XXXX,XX m² / X,XXXX ha
    Perímetro: XXXX,XX m
    """

    # Conversões
    area_fmt = f"{area_m2:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    area_ha = area_m2 / 10000
    area_ha_fmt = f"{area_ha:,.4f}".replace(",", "X").replace(".", ",").replace("X", ".")

    peri_fmt = f"{perimetro_m:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    # Espaço antes
    doc.add_paragraph()

    # -----------------------------
    # TÍTULO DA QUADRA
    # -----------------------------
    p_title = doc.add_paragraph()
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p_title.add_run(f"QUADRA  {str(quadra).zfill(2)}")
    run.bold = True
    run.font.size = Pt(13)

    # -----------------------------
    # ÁREA
    # -----------------------------
    p_area = doc.add_paragraph()
    p_area.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    run_label = p_area.add_run("Área:      ")
    run_label.bold = True
    run_label.font.size = Pt(11)

    run_value = p_area.add_run(f"{area_fmt} m² / {area_ha_fmt} ha")
    run_value.font.size = Pt(11)

    # -----------------------------
    # PERÍMETRO
    # -----------------------------
    p_peri = doc.add_paragraph()
    p_peri.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    run_label = p_peri.add_run("Perímetro: ")
    run_label.bold = True
    run_label.font.size = Pt(11)

    run_value = p_peri.add_run(f"{peri_fmt} m")
    run_value.font.size = Pt(11)

    # Espaço final antes da descrição da quadra
    doc.add_paragraph()
