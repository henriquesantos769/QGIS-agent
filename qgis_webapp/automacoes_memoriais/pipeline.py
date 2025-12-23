import geopandas as gpd
from docx_utils import (_fmt_num_br, _fmt_coord, _add_cabecalho_memorial, fmt_coord,
                         fmt_dist, azimute_dms,add_cabecalho_memorial_quadras, add_bloco_info_quadra)
from docx import Document
from shapely.geometry import LineString, shape, Point, Polygon
import math
from pathlib import Path
from docx.shared import Pt
from qgis.core import (
    QgsApplication, QgsVectorLayer, QgsVectorFileWriter, QgsField,
    QgsProject, QgsCoordinateReferenceSystem, QgsCoordinateTransform,
    QgsCoordinateTransformContext, QgsRasterLayer, QgsWkbTypes
)
from qgis.analysis import QgsNativeAlgorithms
from processing.core.Processing import Processing
import processing
from qgis.PyQt.QtCore import QVariant
from collections import defaultdict
from shapely.ops import unary_union
from shapely.ops import nearest_points
import json
import requests
import numpy as np
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

Processing.initialize()
QgsApplication.processingRegistry().addProvider(QgsNativeAlgorithms())

def _bearing_of_segment(line, ref_pt):
    # pega o segmento mais próximo do ponto de referência e calcula o azimute
    coords = list(line.coords)
    if len(coords) < 2:
        return None
    # escolhe o par (u,v) com menor distância ao ref_pt
    best = None; bestd = 1e18
    for i in range(len(coords)-1):
        seg = gpd.GeoSeries.from_wkt([ ]);  # placeholder to please linters
        p1 = np.array(coords[i]);  p2 = np.array(coords[i+1])
        # ponto médio do segmento
        mid = (p1 + p2)/2.0
        d = (mid[0]-ref_pt.x)**2 + (mid[1]-ref_pt.y)**2
        if d < bestd:
            bestd = d; best = (p1,p2)
    (x1,y1),(x2,y2) = best
    ang = np.degrees(np.arctan2(y2-y1, x2-x1)) % 180.0  # direção de via, sem sentido
    return ang

def atribuir_ruas_e_esquinas_precision(
        upload_dir,
        arquivo_final_nome="final.shp",
        epsg_lotes=31983,
        base_buffer=9,
        min_testada=1.0,
        min_delta_graus=30.0
    ):
    """
    Atribui:
      - Rua: todas as ruas que tocam o lote, em uma string separada por vírgula
      - Esquina: True/False baseado em ângulo das vias e múltiplas testadas
    Salva em final/final_gpkg.gpkg.
    """

    # 1) Carregar dados
    lotes = gpd.read_file(upload_dir / "final" / arquivo_final_nome)
    ruas  = gpd.read_file(upload_dir / "ruas" / "ruas_osm_detalhadas.gpkg")

    # 2) Garantir CRS
    if not lotes.crs:
        lotes.set_crs(epsg=epsg_lotes, inplace=True)

    if not ruas.crs:
        ruas.set_crs(epsg=4326, inplace=True)

    ruas = ruas.to_crs(lotes.crs)

    # 3) Manter somente ruas com nome
    ruas = ruas[ruas["name"].notna()].copy()
    if len(ruas) == 0:
        print("⚠ Nenhuma rua com nome encontrada.")
        lotes["Rua"] = None
        lotes["Esquina"] = False
        out = upload_dir / "final" / "final_gpkg.gpkg"
        lotes.to_file(out, driver="GPKG", encoding="utf-8")
        return out

    # 4) Dissolver ruas por nome + buffer
    ruas_dis = ruas.dissolve(by="name", as_index=False, aggfunc="first")
    ruas_dis["geometry"] = ruas_dis.buffer(base_buffer)
    sidx = ruas_dis.sindex

    # --- helpers internos ---
    def compute_testada(lote_geom, rua_geom):
        borda = lote_geom.boundary
        inter = borda.intersection(rua_geom)
        if inter.is_empty:
            return 0.0
        if inter.geom_type == "LineString":
            return inter.length
        if inter.geom_type == "MultiLineString":
            return sum(g.length for g in inter.geoms)
        return 0.0

    def compute_rua_angle(lote_geom, rua_name):
        eixo = ruas[ruas["name"] == rua_name].union_all()
        ref_pt = nearest_points(lote_geom, eixo)[0]

        if eixo.geom_type == "MultiLineString":
            seg = min(eixo.geoms, key=lambda g: g.distance(lote_geom))
        else:
            seg = eixo

        return _bearing_of_segment(seg, ref_pt)

    # 6) Processar lotes
    ruas_str_final = []
    esquina_final = []

    for _, lote in lotes.iterrows():
        lote_geom = lote.geometry

        # candidatos pelo bbox
        idxs = list(sidx.intersection(lote_geom.bounds))
        cand_ruas = ruas_dis.iloc[idxs]

        touched = []
        angulos = []

        for _, r in cand_ruas.iterrows():
            testada = compute_testada(lote_geom, r.geometry)
            if testada >= min_testada:
                nome = r["name"]
                touched.append(nome)
                ang = compute_rua_angle(lote_geom, nome)
                if ang is not None:
                    angulos.append(ang)

        # nomes únicos ordenados
        touched = sorted(set(touched))

        # Rua: todas as ruas em uma string separada por vírgula
        if len(touched) == 0:
            ruas_str_final.append(None)
        else:
            ruas_str_final.append(", ".join(touched))

        # Esquina: múltiplas ruas com ângulo bem diferente
        if len(angulos) >= 2:
            deltas = []
            for i in range(len(angulos)):
                for j in range(i+1, len(angulos)):
                    d = abs(angulos[i] - angulos[j])
                    d = min(d, 180 - d)
                    deltas.append(d)
            esquina_final.append(max(deltas) >= min_delta_graus)
        else:
            esquina_final.append(False)

    # 7) Guardar nos lotes
    lotes["Rua"] = ruas_str_final
    lotes["Esquina"] = esquina_final

    # 8) Exportar final
    out = upload_dir / "final" / "final_gpkg.gpkg"
    lotes.to_file(out, driver="GPKG", encoding="utf-8")
    print(f"✅ final_gpkg.gpkg gerado com campos Rua e Esquina. Lotes: {len(lotes)}")
    return out


def atribuir_ruas_e_esquinas(upload_dir, arquivo_final_nome="final.shp", buffer_rua=5):
    """
    Atribui a(s) rua(s) correspondente(s) e detecta se cada lote é de esquina.
    Cria um arquivo final.gpkg com as colunas adicionais: 'Rua' e 'Esquina'.

    Parâmetros
    ----------
    upload_dir : Path
        Diretório base do processamento (contém 'ruas/' e 'final/').
    arquivo_final_nome : str
        Nome do shapefile final gerado antes desta etapa (padrão: 'final.shp').
    buffer_rua : float
        Tamanho do buffer em metros aplicado às ruas para detectar contato.
    """
    try:
        print("🏷️ Atribuindo ruas e detectando lotes de esquina...")

        ruas_path = upload_dir / "ruas" / "ruas_osm_detalhadas.gpkg"
        final_path = upload_dir / "final" / arquivo_final_nome
        final_gpkg = upload_dir / "final" / "final_gpkg.gpkg"

        if not ruas_path.exists():
            raise FileNotFoundError(f"Camada de ruas não encontrada: {ruas_path}")
        if not final_path.exists():
            raise FileNotFoundError(f"Camada de lotes não encontrada: {final_path}")

        # Carrega camadas
        gdf_ruas = gpd.read_file(ruas_path)
        gdf_lotes = gpd.read_file(final_path)

        # Garante que ambas estão no mesmo CRS
        if not gdf_lotes.crs:
            print("⚠️ CRS dos lotes indefinido. Definindo como EPSG:31983.")
            gdf_lotes.set_crs(epsg=31983, inplace=True)

        if not gdf_ruas.crs:
            print("⚠️ CRS das ruas indefinido. Definindo como EPSG:4326 (padrão OSM).")
            gdf_ruas.set_crs(epsg=4326, inplace=True)

        # Agora reprojeta corretamente
        gdf_ruas = gdf_ruas.to_crs(gdf_lotes.crs)

        # Buffer pequeno nas ruas (para garantir contato)
        gdf_ruas["geometry"] = gdf_ruas.buffer(buffer_rua)

        # Cria índice espacial
        sindex_ruas = gdf_ruas.sindex

        # Listas para resultados
        ruas_col = []
        esquina_col = []

        for _, lote in gdf_lotes.iterrows():
            # Seleciona ruas próximas (via bounding box)
            possible_idx = list(sindex_ruas.intersection(lote.geometry.bounds))
            possiveis = gdf_ruas.iloc[possible_idx]

            # Filtra ruas que realmente tocam o lote
            ruas_tocadas = possiveis[possiveis.intersects(lote.geometry)]

            # Nomes distintos
            nomes = sorted(set(r.strip() for r in ruas_tocadas["name"].dropna()))

            # Define campos
            ruas_col.append(", ".join(nomes) if nomes else None)
            esquina_col.append(len(nomes) >= 2)

        # Adiciona novas colunas
        gdf_lotes["Rua"] = ruas_col
        gdf_lotes["Rua_lista"] = gdf_lotes["Rua"].apply(lambda x: ";".join(x.split(", ")) if isinstance(x, str) and "," in x else x)
        gdf_lotes["Esquina"] = esquina_col

        # Salva no formato final.gpkg
        final_dir = upload_dir / "final"
        final_dir.mkdir(parents=True, exist_ok=True)
        gdf_lotes.to_file(final_gpkg, driver="GPKG", encoding="utf-8")

        print(f"✅ Atribuição concluída! Arquivo exportado: {final_gpkg}")
        print(f"   Lotes: {len(gdf_lotes)} | Esquinas detectadas: {sum(esquina_col)}")

        return final_gpkg

    except Exception as e:
        print(f"⚠️ Erro na atribuição de ruas e detecção de esquinas: {e}")
        raise

def create_final_gpkg(layer_path: Path) -> Path:
    """
    Cria uma cópia GeoPackage chamada 'final_gpkg.gpkg' a partir da camada shapefile 'final.shp'.
    Retorna o caminho do novo arquivo .gpkg.
    """
    if not layer_path.exists():
        print(f"❌ Arquivo não encontrado: {layer_path}")
        return layer_path

    if layer_path.suffix.lower() != ".shp":
        print(f"⚠️ {layer_path.name} não é um shapefile, ignorando conversão.")
        return layer_path

    gpkg_path = layer_path.parent / "final_gpkg.gpkg"
    print(f"Gerando camada GeoPackage: {gpkg_path.name}")

    layer = QgsVectorLayer(str(layer_path), "final_gpkg", "ogr")
    if not layer.isValid():
        print(f"❌ Falha ao abrir {layer_path.name} para conversão.")
        return layer_path

    # Configura opções de gravação
    options = QgsVectorFileWriter.SaveVectorOptions()
    options.driverName = "GPKG"
    options.layerName = "final_gpkg"
    options.fileEncoding = "UTF-8"

    result = QgsVectorFileWriter.writeAsVectorFormatV3(
        layer,
        str(gpkg_path),
        QgsCoordinateTransformContext(),
        options
    )

    return gpkg_path

def atribuir_ruas_frente(upload_dir,
                         arquivo_final_nome="final.shp",
                         buffer_rua=8,
                         min_testada=1.0):
    """
    Atribui às geometrias de lote:
      - Rua: todas as ruas tocantes (string)
      - Rua_lista: igual à Rua, porém usando ';' para joins
      - Esquina: True/False baseado em múltiplas testadas
      - LadoFrente: rua com maior testada (rua principal)
    """

    print("Atribuindo ruas, esquinas e LadoFrente...")

    ruas_path = upload_dir / "projeto_qgis" / "ruas" / "ruas_osm_detalhadas.gpkg"
    # final_path = upload_dir / "final" / arquivo_final_nome
    final_gpkg = upload_dir / "projeto_qgis" / "final" / "final_gpkg.gpkg"

    if not ruas_path.exists():
        ruas_path = upload_dir / "ruas" / "ruas_osm_detalhadas.gpkg"
    if not final_gpkg.exists():
        final_gpkg = upload_dir / "final" / "final_gpkg.gpkg"

    # ====================
    # 1) Carregar camadas
    # ====================
    gdf_ruas = gpd.read_file(ruas_path)
    gdf_lotes = gpd.read_file(final_gpkg)

    # ====================
    # 2) Garantir CRS
    # ====================
    if not gdf_lotes.crs:
        gdf_lotes.set_crs(epsg=31983, inplace=True)

    if not gdf_ruas.crs:
        gdf_ruas.set_crs(epsg=4326, inplace=True)

    gdf_ruas = gdf_ruas.to_crs(gdf_lotes.crs)

    # ====================
    # 3) Pequeno buffer de contato
    # ====================
    gdf_ruas["geom_buff"] = gdf_ruas.geometry.buffer(buffer_rua)

    # Índice espacial
    sindex_ruas = gdf_ruas.sindex

    # ====================
    # 4) Helpers internos
    # ====================
    def testada(lote_geom, rua_geom):
        inter = lote_geom.boundary.intersection(rua_geom)
        if inter.is_empty:
            return 0.0
        if inter.geom_type == "LineString":
            return inter.length
        if inter.geom_type == "MultiLineString":
            return sum(seg.length for seg in inter.geoms)
        return 0.0

    # ====================
    # 5) Cálculo por lote
    # ====================
    ruas_final = []
    ruas_lista_final = []
    esquina_final = []
    lado_frente_final = []

    for _, lote in gdf_lotes.iterrows():
        lote_geom = lote.geometry

        # ruas candidatas pelo bbox
        idxs = list(sindex_ruas.intersection(lote_geom.bounds))
        cand = gdf_ruas.iloc[idxs]

        # Informação por rua
        ruas_tocantes = []
        testadas = []

        for _, r in cand.iterrows():
            nome = r["name"]
            if not isinstance(nome, str):
                continue

            t = testada(lote_geom, r.geom_buff)
            if t >= min_testada:
                ruas_tocantes.append(nome)
                testadas.append((nome, t))

        # Ordena nomes e salva Rua / Rua_lista
        if ruas_tocantes:
            ruas_ordenadas = sorted(set(ruas_tocantes))
            ruas_final.append(", ".join(ruas_ordenadas))
            ruas_lista_final.append(";".join(ruas_ordenadas))
        else:
            ruas_final.append(None)
            ruas_lista_final.append(None)

        # Esquina = >= 2 ruas tocantes
        esquina_final.append(len(ruas_tocantes) >= 2)

        # -------------------------
        # 🌟 LadoFrente = maior testada
        # -------------------------
        if testadas:
            nome_frente = max(testadas, key=lambda x: x[1])[0]
            lado_frente_final.append(nome_frente)
        else:
            lado_frente_final.append(None)

    # ====================
    # 6) Atualizar GeoDataFrame
    # ====================
    gdf_lotes["Rua"] = ruas_final
    gdf_lotes["Rua_lista"] = ruas_lista_final
    gdf_lotes["Esquina"] = esquina_final
    gdf_lotes["LadoFrente"] = lado_frente_final

    # ====================
    # 7) Exportar
    # ====================
    final_dir = upload_dir / "final"
    final_dir.mkdir(parents=True, exist_ok=True)

    gdf_lotes.to_file(final_gpkg, driver="GPKG", encoding="utf-8")

    print(f"Atribuição concluída: {final_gpkg}")
    print(f"   Lotes: {len(gdf_lotes)}")
    print(f"   Esquinas detectadas: {sum(esquina_final)}")
    print("   Campo LadoFrente criado com sucesso.")

    return final_gpkg

def gerar_confrontacoes(
    upload_dir,
    arquivo_final_nome="final_gpkg.gpkg",
    buffer_rua=8,
    buffer_outros=7,
    epsg_lotes=31983,
    campo_nome_outros="nome"
):
    import math
    import geopandas as gpd
    from shapely.geometry import LineString, Point
    from shapely.geometry.polygon import orient
    from collections import defaultdict

    print("Gerando confrontações (frente geométrica estável)…")

    # ---------------------------------------------------------
    # Parâmetros
    # ---------------------------------------------------------
    TOL = 0.05
    MIN_LEN_LOTE = 1.0
    MIN_FRAC_SEG = 0.20

    PROBE_DIST = 3.0
    MIN_CONTATO_RUA = 1.0
    MIN_FRAC_RUA = 0.30

    # ---------------------------------------------------------
    # Carregar dados
    # ---------------------------------------------------------
    lotes_path = upload_dir / "projeto_qgis" / "final" / arquivo_final_nome
    ruas_path  = upload_dir / "projeto_qgis" / "ruas" / "ruas_osm_detalhadas.gpkg"

    if not lotes_path.exists():
        lotes_path = upload_dir / "final" / arquivo_final_nome
    if not ruas_path.exists():
        ruas_path = upload_dir / "ruas" / "ruas_osm_detalhadas.gpkg"

    gdf_lotes = gpd.read_file(lotes_path)
    gdf_ruas  = gpd.read_file(ruas_path)

    if not gdf_lotes.crs:
        gdf_lotes.set_crs(epsg_lotes, inplace=True)

    gdf_ruas = gdf_ruas.to_crs(gdf_lotes.crs)
    gdf_ruas["geom_buff"] = gdf_ruas.geometry.buffer(buffer_rua)

    sidx_lotes = gdf_lotes.sindex
    sidx_ruas  = gdf_ruas.sindex

    outros_path = upload_dir / "projeto_qgis" / "outros.gpkg"
    gdf_outros = None
    sidx_outros = None

    if outros_path.exists():
        gdf_outros = gpd.read_file(outros_path).to_crs(gdf_lotes.crs)
        gdf_outros["geom_buff"] = gdf_outros.geometry.buffer(buffer_outros)
        sidx_outros = gdf_outros.sindex
    else:
        outros_path = upload_dir / "outros.gpkg"
        if outros_path.exists():
            gdf_outros = gpd.read_file(outros_path).to_crs(gdf_lotes.crs)
            gdf_outros["geom_buff"] = gdf_outros.geometry.buffer(buffer_outros)
            sidx_outros = gdf_outros.sindex

    # ---------------------------------------------------------
    # Garantir campos
    # ---------------------------------------------------------
    for campo in [
        "Conf_Frente",
        "Conf_Direita",
        "Conf_Esquerda",
        "Conf_Fundos",
        "Frente_MX",
        "Frente_MY"
    ]:
        if campo not in gdf_lotes.columns:
            gdf_lotes[campo] = None

    # ---------------------------------------------------------
    # Helpers
    # ---------------------------------------------------------
    def segmentos(coords):
        return [(coords[i], coords[(i + 1) % len(coords)]) for i in range(len(coords))]

    def normal_externa_ccw(p1, p2):
        dx = p2[0] - p1[0]
        dy = p2[1] - p1[1]
        L = math.hypot(dx, dy)
        if L == 0:
            return (0.0, 0.0)
        return (dy / L, -dx / L)

    def idx_por_ponto(coords, mx, my):
        P = Point(mx, my)
        best_i, best_d = None, 1e30
        for i in range(len(coords)):
            seg = LineString([coords[i], coords[(i + 1) % len(coords)]])
            d = seg.distance(P)
            if d < best_d:
                best_d, best_i = d, i
        return best_i

    def soma_confronto(seg_infos, indices):
        soma = defaultdict(float)

        for s in seg_infos:
            if s["idx"] in indices:
                soma[s["nome"]] += s["len"]

        if not soma:
            return "Área não identificada"

        # 🔥 prioridade: qualquer nome válido vence
        for nome in soma:
            if nome != "Área não identificada":
                return nome

        return "Área não identificada"

    # ---------------------------------------------------------
    # PROCESSAMENTO PRINCIPAL
    # ---------------------------------------------------------
    for idx, row in gdf_lotes.iterrows():

        geom = row.geometry
        if geom is None or geom.is_empty:
            continue

        lado_frente = row.get("LadoFrente")
        if not isinstance(lado_frente, str) or not lado_frente.strip():
            continue

        geom = orient(geom, sign=-1.0)
        coords = list(geom.exterior.coords)
        if coords[0] == coords[-1]:
            coords = coords[:-1]

        if len(coords) < 3:
            continue

        segs = segmentos(coords)
        seg_infos = []

        # ---------------------------
        # 1) Confrontante por segmento
        # ---------------------------
        for i, seg in enumerate(segs):
            ln = LineString(seg)
            seg_len = ln.length
            if seg_len == 0:
                continue

            ln_buff = ln.buffer(TOL)
            melhor_nome = "Área não identificada"
            melhor_score = 0.0

            # Lotes vizinhos
            for il in sidx_lotes.intersection(ln.bounds):
                if il == idx:
                    continue
                other = gdf_lotes.iloc[il]
                inter = other.geometry.boundary.intersection(ln)

                if inter.is_empty:
                    continue

                if inter.geom_type not in ("LineString", "MultiLineString"):
                    continue


                if not inter.is_empty:
                    L = inter.length
                    if not (
                            L >= MIN_LEN_LOTE or
                            (L / seg_len) >= MIN_FRAC_SEG
                        ):
                            continue
                    if L > melhor_score:
                        melhor_score = L
                        melhor_nome = f"Lote {other.get('lote_num')}"

            # Ruas
            if melhor_score == 0:
                for ir in sidx_ruas.intersection(ln.bounds):
                    rua = gdf_ruas.iloc[ir]
                    nome_rua = rua.get("name")
                    if not isinstance(nome_rua, str):
                        continue
                    inter = rua["geom_buff"].intersection(ln)
                    if not inter.is_empty:
                        L = inter.length
                        if L >= MIN_CONTATO_RUA and (L / seg_len) >= MIN_FRAC_RUA and L > melhor_score:
                            melhor_score = L
                            melhor_nome = nome_rua

            # Outros
            if gdf_outros is not None:
                for io in sidx_outros.intersection(ln.bounds):
                    outro = gdf_outros.iloc[io]
                    nome_outro = outro.get(campo_nome_outros)

                    if not isinstance(nome_outro, str) or not nome_outro.strip():
                        continue

                    inter = outro["geom_buff"].intersection(ln)
                    if inter.is_empty:
                        continue

                    # importante: filtrar toque em vértice
                    if inter.geom_type not in ("LineString", "MultiLineString"):
                        continue

                    L = inter.length

                    if (L / seg_len) < MIN_FRAC_SEG:
                        continue

                    if L > melhor_score:
                        melhor_score = L
                        melhor_nome = nome_outro

            seg_infos.append({
                "idx": i,
                "seg": seg,
                "len": seg_len,
                "nome": melhor_nome
            })

        # ---------------------------
        # 2) Escolher FRENTE geométrica
        # ---------------------------
        ruas_frente = gdf_ruas[gdf_ruas["name"] == lado_frente]
        if ruas_frente.empty:
            continue

        candidatos_fortes = []
        candidatos_fracos = []

        for s in seg_infos:
            if s["nome"] != lado_frente:
                continue

            p1, p2 = s["seg"]
            ln = LineString([p1, p2])
            seg_len = s["len"]

            contato = 0.0
            for _, rua in ruas_frente.iterrows():
                inter = rua["geom_buff"].intersection(ln)
                if not inter.is_empty:
                    contato = max(contato, inter.length)

            if contato <= 0:
                continue

            mx = (p1[0] + p2[0]) / 2
            my = (p1[1] + p2[1]) / 2
            nx, ny = normal_externa_ccw(p1, p2)
            probe = Point(mx + nx * PROBE_DIST, my + ny * PROBE_DIST)

            hit = any(rua["geom_buff"].contains(probe) for _, rua in ruas_frente.iterrows())

            cand = {
                "idx": s["idx"],
                "mx": mx,
                "my": my,
                "contato": contato,
                "seg_len": seg_len
            }

            if hit:
                candidatos_fortes.append(cand)
            else:
                candidatos_fracos.append(cand)

        if candidatos_fortes:
            frente = max(candidatos_fortes, key=lambda c: (c["contato"], c["seg_len"]))
        elif candidatos_fracos:
            frente = max(candidatos_fracos, key=lambda c: (c["contato"], c["seg_len"]))
        else:
            candidatos_brutos = []

            for s in seg_infos:
                p1, p2 = s["seg"]
                ln = LineString([p1, p2])

                contato = 0.0
                for _, rua in ruas_frente.iterrows():
                    inter = rua["geom_buff"].intersection(ln)
                    if not inter.is_empty:
                        contato = max(contato, inter.length)

                if contato > 0:
                    mx = (p1[0] + p2[0]) / 2
                    my = (p1[1] + p2[1]) / 2
                    candidatos_brutos.append({
                        "idx": s["idx"],
                        "mx": mx,
                        "my": my,
                        "contato": contato,
                        "seg_len": s["len"]
                    })

            if not candidatos_brutos:
                continue

            frente = max(candidatos_brutos, key=lambda c: (c["contato"], c["seg_len"]))

        gdf_lotes.at[idx, "Frente_MX"] = frente["mx"]
        gdf_lotes.at[idx, "Frente_MY"] = frente["my"]

        # ---------------------------
        # 3) Confrontações por lado
        # ---------------------------
        idx_frente = idx_por_ponto(coords, frente["mx"], frente["my"])
        lados = classificar_lados_por_frente(coords, idx_frente)

        gdf_lotes.at[idx, "Conf_Frente"]   = soma_confronto(seg_infos, lados["frente"])
        gdf_lotes.at[idx, "Conf_Direita"]  = soma_confronto(seg_infos, lados["direita"])
        gdf_lotes.at[idx, "Conf_Esquerda"] = soma_confronto(seg_infos, lados["esquerda"])
        gdf_lotes.at[idx, "Conf_Fundos"]   = soma_confronto(seg_infos, lados["fundos"])

    # ---------------------------------------------------------
    # Salvar resultado
    # ---------------------------------------------------------
    out = upload_dir / "final" / "final_confrontacoes.gpkg"
    gdf_lotes.to_file(out, driver="GPKG", encoding="utf-8")

    print("Confrontações geradas com sucesso (sem debug, frente correta).")
    print("Arquivo:", out)
    return out

def calcular_medidas_e_azimutes(
    upload_dir,
    arquivo_final_nome="final_confrontacoes.gpkg",
    epsg_lotes=31983
):
    """
    Calcula para cada lote:
      - Comprimentos oficiais (frente, fundos, direita, esquerda)
      - Azimutes (graus decimais)
      - Rumos (formato N xx°xx'xx" E)

    🔒 Usa Frente_MX / Frente_MY como referência absoluta da frente.
    """

    import math
    import geopandas as gpd
    from shapely.geometry import LineString, Point
    from shapely.geometry.polygon import orient

    print("Calculando medidas e azimutes (baseado na frente geométrica)…")

    # --------------------------------------------------
    # Carregar dados
    # --------------------------------------------------
    lotes_path = upload_dir / "final" / arquivo_final_nome
    gdf = gpd.read_file(lotes_path)

    if not gdf.crs:
        gdf.set_crs(epsg_lotes, inplace=True)

    # --------------------------------------------------
    # Helpers
    # --------------------------------------------------
    def segmentos(coords):
        return [(coords[i], coords[(i + 1) % len(coords)]) for i in range(len(coords))]

    def azimute(seg):
        (x1, y1), (x2, y2) = seg
        ang = math.degrees(math.atan2(y2 - y1, x2 - x1))
        return (ang + 360) % 360

    def rumo_from_az(az):
        if 0 <= az < 90:
            q1, q2, ang = "N", "E", az
        elif 90 <= az < 180:
            q1, q2, ang = "S", "E", 180 - az
        elif 180 <= az < 270:
            q1, q2, ang = "S", "W", az - 180
        else:
            q1, q2, ang = "N", "W", 360 - az

        g = int(ang)
        m_float = (ang - g) * 60
        m = int(m_float)
        s = (m_float - m) * 60

        return f"{q1} {g:02d}°{m:02d}'{s:04.1f}\" {q2}"

    def idx_segmento_por_ponto(coords, mx, my):
        """Retorna o índice do segmento mais próximo ao ponto (mx,my)."""
        P = Point(mx, my)
        best_i, best_d = None, 1e18
        for i in range(len(coords)):
            seg = LineString([coords[i], coords[(i + 1) % len(coords)]])
            d = seg.distance(P)
            if d < best_d:
                best_d, best_i = d, i
        return best_i

    # --------------------------------------------------
    # Campos a criar
    # --------------------------------------------------
    lados = ["Frente", "Fundos", "Direita", "Esquerda"]

    for lado in lados:
        for campo in [f"Comp_{lado}", f"Az_{lado}", f"Rumo_{lado}"]:
            if campo not in gdf.columns:
                gdf[campo] = None

    # --------------------------------------------------
    # LOOP lote a lote
    # --------------------------------------------------
    for idx, row in gdf.iterrows():

        geom = row.geometry
        if geom is None or geom.is_empty:
            continue

        mx = row.get("Frente_MX")
        my = row.get("Frente_MY")

        if mx is None or my is None:
            continue

        geom = orient(geom, sign=-1.0)
        coords = list(geom.exterior.coords)
        if coords[0] == coords[-1]:
            coords = coords[:-1]

        n = len(coords)
        if n < 3:
            continue

        # 🔒 índice da frente recalculado de forma robusta
        idx_frente = idx_segmento_por_ponto(coords, mx, my)

        if idx_frente is None:
            continue

        # classificar lados
        lados_geom = classificar_lados_por_frente(coords, idx_frente)

        segs = segmentos(coords)

        # --------------------------------------------------
        # Para cada lado: soma comprimentos e define az/rumo
        # --------------------------------------------------
        for nome_lado, indices in [
            ("Frente",   lados_geom["frente"]),
            ("Fundos",   lados_geom["fundos"]),
            ("Direita",  lados_geom["direita"]),
            ("Esquerda", lados_geom["esquerda"]),
        ]:

            if not indices:
                continue

            comp_total = 0.0
            seg_dominante = None
            maior_len = -1.0

            for i in indices:
                seg = segs[i]
                L = LineString(seg).length
                comp_total += L

                if L > maior_len:
                    maior_len = L
                    seg_dominante = seg

            if seg_dominante is None:
                continue

            az = azimute(seg_dominante)
            rumo = rumo_from_az(az)

            gdf.at[idx, f"Comp_{nome_lado}"] = round(comp_total, 2)
            gdf.at[idx, f"Az_{nome_lado}"]   = round(az, 6)
            gdf.at[idx, f"Rumo_{nome_lado}"] = rumo

    # --------------------------------------------------
    # Salvar resultado
    # --------------------------------------------------
    out = upload_dir / "final" / "final_medidas_azimutes.gpkg"
    gdf.to_file(out, driver="GPKG", encoding="utf-8")

    print("Medidas e azimutes calculados com sucesso (frente geométrica).")
    print("Arquivo:", out)
    return out


def classificar_lados_por_frente(coords, idx_frente):
    """
    Classifica segmentos (índices i referentes a arestas coords[i] -> coords[i+1]) em:
    - frente
    - fundos
    - direita
    - esquerda

    Definição geométrica robusta:
    - "Olhar para a rua" = direção do ponto interno (C) para o meio da frente.
    - direita/esquerda definidas por essa direção (posição), não pela orientação do segmento.
    - independe de CW/CCW e independe do sentido (f1->f2) do segmento de frente.
    """
    import math
    from shapely.geometry import Polygon

    n = len(coords)
    if n < 3:
        return {"frente": [idx_frente], "fundos": [], "direita": [], "esquerda": []}

    poly = Polygon(coords)
    C = poly.representative_point()  # ponto garantidamente interno

    # --- Meio da frente ---
    x1, y1 = coords[idx_frente]
    x2, y2 = coords[(idx_frente + 1) % n]
    mx_f = (x1 + x2) / 2.0
    my_f = (y1 + y2) / 2.0

    # --- Eixo "forward": de dentro (C) apontando para a frente/rua ---
    fx = mx_f - C.x
    fy = my_f - C.y
    Lf = math.hypot(fx, fy)
    if Lf == 0:
        return {"frente": [idx_frente], "fundos": [], "direita": [], "esquerda": []}

    fx /= Lf
    fy /= Lf

    # --- Eixo "right": rotação de +90° do forward ---
    rx, ry = (fy, -fx)

    # --- Para fundos: projeção no eixo forward (frente = alto, fundos = baixo) ---
    proj_f = []
    for i in range(n):
        if i == idx_frente:
            continue
        ax, ay = coords[i]
        bx, by = coords[(i + 1) % n]
        mx = (ax + bx) / 2.0 - C.x
        my = (ay + by) / 2.0 - C.y
        depth = mx * fx + my * fy
        proj_f.append((i, depth))

    if not proj_f:
        return {"frente": [idx_frente], "fundos": [], "direita": [], "esquerda": []}

    depths = [d for _, d in proj_f]
    min_depth = min(depths)  # mais negativo = mais "pra trás" = fundos

    frente_list = [idx_frente]
    fundos_list = []
    direita_list = []
    esquerda_list = []

    for i, depth in proj_f:
        # FUNDOS: bem próximo do "mais fundo" real (critério restritivo)
        # (se quiser, ajuste 0.15 -> 0.20 conforme seus lotes)
        if depth <= min_depth * 0.85:  # min_depth é negativo; 0.85 mantém perto do extremo
            fundos_list.append(i)
            continue

        # Direita/esquerda: projeção no eixo "right"
        ax, ay = coords[i]
        bx, by = coords[(i + 1) % n]
        mx = (ax + bx) / 2.0 - C.x
        my = (ay + by) / 2.0 - C.y
        side = mx * rx + my * ry

        if side > 0:
            direita_list.append(i)
        else:
            esquerda_list.append(i)

    return {
        "frente": frente_list,
        "fundos": fundos_list,
        "direita": direita_list,
        "esquerda": esquerda_list
    }

def _memorial_lote_completo(
    row,
    nucleo,
    municipio,
    uf,
    gdf_lotes,
    gdf_ruas,
    gdf_outros=None
):
    from shapely.geometry import Point, LineString
    from shapely.geometry.polygon import orient

    geom = row.geometry
    if geom is None or geom.is_empty:
        return "Geometria indisponível."

    quadra = row.get("quadra")
    lote   = row.get("lote_num")

    geom = orient(geom, sign=-1.0)
    coords = list(geom.exterior.coords)
    if coords[0] == coords[-1]:
        coords = coords[:-1]

    n = len(coords)
    if n < 3:
        return "Lote com geometria insuficiente."

    mx = row.get("Frente_MX")
    my = row.get("Frente_MY")
    if mx is None or my is None:
        return "Frente geométrica não definida."

    P_frente = Point(mx, my)

    def idx_segmento_por_ponto(coords, P):
        best_i, best_d = None, 1e18
        for i in range(len(coords)):
            seg = LineString([coords[i], coords[(i + 1) % n]])
            d = seg.distance(P)
            if d < best_d:
                best_d, best_i = d, i
        return best_i

    idx_frente = idx_segmento_por_ponto(coords, P_frente) or 0
    lados = classificar_lados_por_frente(coords, idx_frente)

    # parâmetros (iguais à sua lógica)
    MIN_LEN_LOTE = 1.0
    MIN_FRAC_SEG = 0.20

    MIN_CONTATO_RUA = 1.0
    MIN_FRAC_RUA = 0.30

    def confronto_por_segmento(i):
        # fallback se não tiver datasets carregados
        if gdf_lotes is None or gdf_ruas is None:
            return "Área não identificada"

        p1 = coords[i]
        p2 = coords[(i + 1) % n]
        ln = LineString([p1, p2])
        seg_len = ln.length
        if seg_len == 0:
            return "Área não identificada"

        melhor_nome = "Área não identificada"
        melhor_score = 0.0

        # 1) lotes vizinhos
        # IMPORTANTE: não usar "sindex" do row; use do gdf (injete também se quiser performance)
        for _, other in gdf_lotes.iterrows():
            if other.get("lote_num") == lote:
                continue
            inter = other.geometry.boundary.intersection(ln)
            if inter.is_empty:
                continue
            if inter.geom_type not in ("LineString", "MultiLineString"):
                continue
            L = inter.length

            # regra híbrida: absoluto OU relativo (para segmentos < 1m)
            if not (L >= MIN_LEN_LOTE or (L / seg_len) >= MIN_FRAC_SEG):
                continue

            if L > melhor_score:
                melhor_score = L
                melhor_nome = f"Lote {other.get('lote_num')}"

        # 2) ruas (usa buffer das ruas, como no teu gerador)
        if melhor_score == 0.0 and "geom_buff" in gdf_ruas.columns:
            for _, rua in gdf_ruas.iterrows():
                nome_rua = rua.get("name")
                if not isinstance(nome_rua, str) or not nome_rua.strip():
                    continue
                inter = rua["geom_buff"].intersection(ln)
                if inter.is_empty:
                    continue
                if inter.geom_type not in ("LineString", "MultiLineString"):
                    continue
                L = inter.length
                if L >= MIN_CONTATO_RUA and (L / seg_len) >= MIN_FRAC_RUA and L > melhor_score:
                    melhor_score = L
                    melhor_nome = nome_rua

        # 3) outros
        if gdf_outros is not None and "geom_buff" in gdf_outros.columns:
            for _, outro in gdf_outros.iterrows():
                nome_outro = outro.get("nome")
                if not isinstance(nome_outro, str) or not nome_outro.strip():
                    continue
                inter = outro["geom_buff"].intersection(ln)
                if inter.is_empty:
                    continue
                if inter.geom_type not in ("LineString", "MultiLineString"):
                    continue
                L = inter.length
                if (L / seg_len) >= MIN_FRAC_SEG and L > melhor_score:
                    melhor_score = L
                    melhor_nome = nome_outro

        return melhor_nome

    # Helpers de texto
    def fmt_coord(v):
        return format(float(v), ",.4f").replace(",", "X").replace(".", ",").replace("X", ".")

    def fmt_dist(v):
        return format(float(v), ",.2f").replace(",", "X").replace(".", ",").replace("X", ".")

    def nome_lado(i):
        if i in lados["frente"]:
            return "de frente"
        if i in lados["fundos"]:
            return "ao fundo"
        if i in lados["direita"]:
            return "do lado direito"
        if i in lados["esquerda"]:
            return "do lado esquerdo"
        return "pelo perímetro"

    def deflexao(k):
        if k == 0:
            return ""
        i_prev = (idx_frente + k - 1) % n
        i_curr = (idx_frente + k) % n
        i_next = (idx_frente + k + 1) % n

        x1, y1 = coords[i_prev]
        x2, y2 = coords[i_curr]
        x3, y3 = coords[i_next]

        v1 = (x2 - x1, y2 - y1)
        v2 = (x3 - x2, y3 - y2)

        cross = v1[0] * v2[1] - v1[1] * v2[0]
        if cross < 0:
            return "deste ponto deflete à direita"
        elif cross > 0:
            return "deste ponto deflete à esquerda"
        else:
            return "deste ponto segue em linha"

    area = geom.area
    perim = geom.length

    intro = (
        f"O lote de terreno sob nº {lote} da Quadra {quadra}, do Núcleo denominado "
        f"“{nucleo}”, no município de {municipio} - {uf}, de formato irregular, "
        f"abrangendo uma área de {fmt_coord(area)} m² e um perímetro de {fmt_coord(perim)} m. "
    )

    partes = [intro]

    for k in range(n):
        i = (idx_frente + k) % n
        x1, y1 = coords[i]
        x2, y2 = coords[(i + 1) % n]

        dist = Point(x1, y1).distance(Point(x2, y2))
        dist_fmt = fmt_dist(dist)

        coord1 = f"(EX: {fmt_coord(x1)}  NY: {fmt_coord(y1)})"
        coord2 = f"(EX: {fmt_coord(x2)}  NY: {fmt_coord(y2)})"

        confronto = confronto_por_segmento(i)  # ✅ aqui muda tudo
        tipo_lado = nome_lado(i)
        frase_deflexao = deflexao(k)

        if k == 0:
            texto = (
                f"Para quem de dentro do lote {lote} olha para {confronto}, "
                f"inicia-se a descrição na coordenada {coord1}, "
                f"com uma distância de {dist_fmt} m {tipo_lado} "
                f"até a coordenada {coord2}, confrontando com {confronto}, "
            )
        elif k < n - 1:
            texto = (
                f"{frase_deflexao} com uma distância de {dist_fmt} m {tipo_lado} "
                f"até a coordenada {coord2}, confrontando com {confronto}, "
            )
        else:
            texto = (
                f"{frase_deflexao} com uma distância de {dist_fmt} m {tipo_lado} "
                f"até a coordenada {coord2}, confrontando com {confronto};"
            )

        partes.append(texto)

    return " ".join(partes)

def gerar_memorial_quadra(
    upload_dir: Path,
    arquivo_final_nome: str,
    quadra_alvo,
    nucleo: str,
    municipio: str,
    uf: str,
    promotor: str = "Instituto Cidade Legal",
    saida_dir: Path | None = None,
    buffer_outros: float = 7.0,
):

    # --------------------------------------------------
    # Caminhos
    # --------------------------------------------------
    final_path = upload_dir / "final" / arquivo_final_nome
    ruas_path  = upload_dir / "ruas" / "ruas_osm_detalhadas.gpkg"

    outros_path1 = upload_dir / "projeto_qgis" / "outros.gpkg"
    outros_path2 = upload_dir / "outros.gpkg"

    # --------------------------------------------------
    # Carregar dados
    # --------------------------------------------------
    gdf_lotes = gpd.read_file(final_path)
    gdf_ruas  = gpd.read_file(ruas_path).to_crs(gdf_lotes.crs)

    # buffer das ruas (necessário para o memorial por segmento)
    if "geom_buff" not in gdf_ruas.columns:
        gdf_ruas["geom_buff"] = gdf_ruas.geometry.buffer(7)

    # OUTROS (opcional)
    gdf_outros = None
    if outros_path1.exists():
        gdf_outros = gpd.read_file(outros_path1).to_crs(gdf_lotes.crs)
    elif outros_path2.exists():
        gdf_outros = gpd.read_file(outros_path2).to_crs(gdf_lotes.crs)

    if gdf_outros is not None:
        gdf_outros["geom_buff"] = gdf_outros.geometry.buffer(buffer_outros)

    # --------------------------------------------------
    # Filtrar quadra
    # --------------------------------------------------
    gdf_lotes["quadra_str"] = gdf_lotes["quadra"].astype(str)
    quadra_str = str(quadra_alvo)

    lotes_quadra = gdf_lotes[gdf_lotes["quadra_str"] == quadra_str].copy()
    if lotes_quadra.empty:
        print(f"⚠ Nenhum lote encontrado para a quadra {quadra_str}.")
        return None

    lotes_quadra = lotes_quadra.sort_values(by="lote_num")

    # --------------------------------------------------
    # Saída
    # --------------------------------------------------
    if saida_dir is None:
        saida_dir = upload_dir / "memoriais"
    saida_dir.mkdir(parents=True, exist_ok=True)

    doc_name = f"memorial_quadra_{quadra_str}.docx"
    docx_path = saida_dir / doc_name

    doc = Document()

    # --------------------------------------------------
    # Cabeçalho
    # --------------------------------------------------
    _add_cabecalho_memorial(
        doc,
        titulo="MEMORIAL DESCRITIVO",
        quadra=quadra_str,
        nucleo=nucleo,
        municipio=municipio,
        uf=uf,
        promotor=promotor
    )

    # --------------------------------------------------
    # Percorrer lotes
    # --------------------------------------------------
    for _, row in lotes_quadra.iterrows():
        quadra = row.get("quadra")
        lote_num = row.get("lote_num")
        geom = row.geometry

        area_m2 = geom.area if geom is not None else None
        perimetro_m = geom.length if geom is not None else None

        # Título do lote
        p_header_lote = doc.add_paragraph()
        run = p_header_lote.add_run(f"Quadra: {quadra}\nLote: {lote_num}")
        run.bold = True

        if area_m2 is not None:
            run2 = p_header_lote.add_run(f"\nÁrea: {_fmt_num_br(area_m2, 2)} m²\n")
            run2.bold = True

        doc.add_paragraph()

        # 🔥 CHAMADA CORRETA DO MEMORIAL
        texto_lote = _memorial_lote_completo(
            row=row,
            nucleo=nucleo,
            municipio=municipio,
            uf=uf,
            gdf_lotes=gdf_lotes,
            gdf_ruas=gdf_ruas,
            gdf_outros=gdf_outros
        )

        p_desc_lote = doc.add_paragraph()
        p_desc_lote.paragraph_format.first_line_indent = Pt(12)
        p_desc_lote.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_desc_lote.add_run(texto_lote)

        doc.add_paragraph()

    # --------------------------------------------------
    # Rodapé
    # --------------------------------------------------
    doc.add_paragraph()
    rod = doc.add_paragraph()
    rod.add_run(
        "Todas as medidas lineares, áreas e coordenadas foram calculadas no sistema de projeção "
        "e datum adotados no projeto (ex.: SIRGAS2000 / UTM)."
    )

    doc.save(str(docx_path))
    print(f"Memorial da quadra {quadra_str} salvo em: {docx_path}")
    return docx_path

def gerar_memoriais_em_lote(upload_dir: Path,
                            arquivo_final_nome: str = "final_medidas_azimutes.gpkg",
                            nucleo: str = "Teste",
                            municipio: str = "Teste",
                            uf: str = "Teste",
                            promotor: str = "Instituto Cidade Legal",
                            saida_dir: Path | None = None):

    final_path = upload_dir / "final" / arquivo_final_nome
    gdf = gpd.read_file(final_path)

    if "quadra" not in gdf.columns:
        raise ValueError("Coluna 'quadra' não encontrada no GPKG final.")

    if saida_dir is None:
        saida_dir = upload_dir / "memoriais"
    saida_dir.mkdir(parents=True, exist_ok=True)

    quadras_unicas = sorted(gdf["quadra"].unique(), key=lambda x: str(x))

    paths_gerados = []

    for q in quadras_unicas:
        path_q = gerar_memorial_quadra(
            upload_dir=upload_dir,
            arquivo_final_nome=arquivo_final_nome,
            quadra_alvo=q,
            nucleo=nucleo,
            municipio=municipio,
            uf=uf,
            promotor=promotor,
            saida_dir=saida_dir
        )
        if path_q:
            paths_gerados.append(path_q)

    print("Geração de memoriais concluída.")
    print("Arquivos gerados:")
    for p in paths_gerados:
        print("  -", p)

    return paths_gerados

def gerar_geometrias_quadras(
    upload_dir: Path,
    arquivo_final_nome: str = "final_medidas_azimutes.gpkg",
    epsg_lotes: int = 31983,
    saida_nome: str = "quadras_contorno.gpkg"
) -> Path:
    """
    Gera o contorno de cada quadra a partir dos lotes,
    unindo os polígonos e calculando área e perímetro da quadra.

    Saída: GPKG em upload_dir / "final" / saida_nome
    com colunas:
      - quadra
      - geometry (polígono da quadra)
      - area_q  (área da quadra)
      - perim_q (perímetro da quadra)
    """
    final_path = upload_dir / "final" / arquivo_final_nome
    if not final_path.exists():
        raise FileNotFoundError(f"Arquivo de lotes não encontrado: {final_path}")

    gdf_lotes = gpd.read_file(final_path)

    if "quadra" not in gdf_lotes.columns:
        raise ValueError("Coluna 'quadra' não encontrada no arquivo de lotes.")

    # Garantir CRS
    if not gdf_lotes.crs:
        gdf_lotes.set_crs(epsg=epsg_lotes, inplace=True)

    quadra_records = []

    # Agrupa lotes por quadra
    for quadra_val, sub in gdf_lotes.groupby("quadra"):
        geoms = [g for g in sub.geometry if g is not None]

        if not geoms:
            continue

        # União das geometrias da quadra
        uni = unary_union(geoms)

        # Limpar geometrias zoadas
        uni = uni.buffer(0)

        # Se ainda for MultiPolygon, pega o maior polígono como contorno principal
        if uni.geom_type == "MultiPolygon":
            uni = max(uni.geoms, key=lambda g: g.area)

        # Se ainda for GeometryCollection / outra coisa, tentar convex hull
        if not isinstance(uni, Polygon):
            uni = uni.convex_hull

        area_q = uni.area
        perim_q = uni.length

        quadra_records.append({
            "quadra": quadra_val,
            "geometry": uni,
            "area_q": area_q,
            "perim_q": perim_q,
        })

    if not quadra_records:
        raise RuntimeError("Não foi possível gerar geometrias de quadra (lista vazia).")

    gdf_quadras = gpd.GeoDataFrame(quadra_records, geometry="geometry", crs=gdf_lotes.crs)

    out_path = upload_dir / "final" / saida_nome
    out_path.parent.mkdir(parents=True, exist_ok=True)
    gdf_quadras.to_file(out_path, driver="GPKG", encoding="utf-8")

    print(f"Geometrias de quadra geradas: {out_path} (quadras: {len(gdf_quadras)})")
    return out_path

def segmentar_quadra_com_confrontantes(
    upload_dir: Path,
    quadras_gpkg: str = "quadras_contorno.gpkg",
    lotes_gpkg: str = "final_medidas_azimutes.gpkg",
    ruas_gpkg: str = "ruas_osm_detalhadas.gpkg",
    buffer_rua: float = 9.0
):
    """
    Segmenta cada quadra em trechos individuais e determina
    o confrontante de cada segmento: rua, lote ou limite.

    Retorna um GeoDataFrame com:
      quadra, seq, geometry, azimute, comprimento,
      x1, y1, x2, y2,
      tipo ('rua' | 'lote' | 'limite'),
      confronto (nome da rua ou 'Lote xx' ou 'Limite')
    """

    # ----------------------------------------------------------
    # Carregar camadas
    # ----------------------------------------------------------
    quadras_path = upload_dir / "final" / quadras_gpkg
    lotes_path   = upload_dir / "final" / lotes_gpkg
    ruas_path    = upload_dir / "ruas" / ruas_gpkg

    gdf_quadras = gpd.read_file(quadras_path)
    gdf_lotes   = gpd.read_file(lotes_path)
    gdf_ruas    = gpd.read_file(ruas_path)

    crs = gdf_quadras.crs
    gdf_lotes = gdf_lotes.to_crs(crs)
    gdf_ruas  = gdf_ruas.to_crs(crs)

    # buffer das ruas para facilitar interseção
    gdf_ruas["geom_buff"] = gdf_ruas.geometry.buffer(buffer_rua)
    sidx_ruas  = gdf_ruas.sindex
    sidx_lotes = gdf_lotes.sindex

    registros = []

    # ----------------------------------------------------------
    # Processar quadra por quadra
    # ----------------------------------------------------------
    for _, qrow in gdf_quadras.iterrows():
        quadra = qrow["quadra"]
        geom = qrow.geometry

        coords = list(geom.exterior.coords)

        # cada vértice vira um segmento p1->p2
        segmentos = []
        for i in range(len(coords) - 1):
            x1,y1 = coords[i]
            x2,y2 = coords[i+1]

            line = LineString([(x1,y1),(x2,y2)])
            dx = x2 - x1
            dy = y2 - y1
            az = (math.degrees(math.atan2(dy, dx)) + 360) % 360
            dist = line.length

            segmentos.append({
                "quadra": quadra,
                "seq": i + 1,
                "geometry": line,
                "azimute": az,
                "comprimento": dist,
                "x1": x1, "y1": y1,
                "x2": x2, "y2": y2
            })

        # ----------------------------------------------------------
        # Determinar confrontantes
        # ----------------------------------------------------------
        for seg in segmentos:
            line = seg["geometry"]

            # — Rua?
            rua_nome = None
            bbox = list(sidx_ruas.intersection(line.bounds))

            for idx_r in bbox:
                rr = gdf_ruas.iloc[idx_r]
                if rr["geom_buff"].intersects(line):
                    if rr.get("name"):
                        rua_nome = rr["name"]
                        break

            # — Lote?
            lote_conf = None
            if rua_nome is None:
                bbox2 = list(sidx_lotes.intersection(line.bounds))
                for idx_l in bbox2:
                    lote = gdf_lotes.iloc[idx_l]
                    if lote.geometry.intersects(line):
                        lote_conf = f"Lote {lote.get('lote_num')} - Quadra {lote.get('quadra')}"
                        break

            if rua_nome:
                seg["tipo"] = "rua"
                seg["confronto"] = rua_nome
            #elif lote_conf:
            #    seg["tipo"] = "lote"
            #    seg["confronto"] = lote_conf
            #else:
            #    seg["tipo"] = "limite"
            #    seg["confronto"] = "Limite"

                registros.append(seg)

    gdf = gpd.GeoDataFrame(registros, geometry="geometry", crs=crs)

    # salvar
    out_path = upload_dir / "final" / "quadras_segmentos.gpkg"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    gdf.to_file(out_path, driver="GPKG", encoding="utf-8")

    print(f"✅ Segmentos de quadras gerados: {out_path}")
    return gdf

def calcular_deflexoes_segmentos(
    upload_dir: Path,
    seg_gpkg: str = "quadras_segmentos.gpkg"
):
    """
    Lê os segmentos de quadras e adiciona:
      - deflexão entre trechos ('direita' ou 'esquerda')
      - tipo_lado (frente, direita, esquerda, fundos, perímetro)
      - delta (valor do giro)
      - flag é_primeiro (para o trecho inicial)
    """

    path = upload_dir / "final" / seg_gpkg
    gdf = gpd.read_file(path)

    registros_final = []

    for quadra in sorted(gdf["quadra"].unique(), key=lambda x: str(x)):
        sub = gdf[gdf["quadra"] == quadra].copy()
        sub = sub.sort_values("seq").reset_index(drop=True)

        N = len(sub)

        # percorre todos os segmentos
        for i in range(N):
            row = sub.loc[i]

            az1 = row["azimute"]
            comprimento = row["comprimento"]
            tipo_conf = row["tipo"]
            nome_conf = row["confronto"]

            # ----------------------------------------------------
            # calcular próximo azimute (para deflexão)
            # ----------------------------------------------------
            if i < N - 1:
                az2 = sub.loc[i+1, "azimute"]
            else:
                az2 = sub.loc[0, "azimute"]  # volta ao início

            delta = (az2 - az1 + 360) % 360

            if 0 < delta < 180:
                deflex = "à direita"
            else:
                deflex = "à esquerda"

            # ----------------------------------------------------
            # tipo de lado — baseado no confrontante
            # ----------------------------------------------------
            if tipo_conf == "rua":
                tipo_lado = "de frente"
            elif tipo_conf == "lote":
                tipo_lado = "do lado direito"  # pode melhorar depois
            elif tipo_conf == "limite":
                tipo_lado = "pelo perímetro"
            else:
                tipo_lado = ""

            registros_final.append({
                "quadra": quadra,
                "seq": row["seq"],
                "geometry": row.geometry,
                "x1": row["x1"],
                "y1": row["y1"],
                "x2": row["x2"],
                "y2": row["y2"],
                "comprimento": comprimento,
                "azimute": az1,
                "delta": delta,
                "deflexao": deflex,
                "tipo": tipo_conf,
                "tipo_lado": tipo_lado,
                "confronto": nome_conf,
                "primeiro": (i == 0)
            })

    gdf_out = gpd.GeoDataFrame(registros_final, geometry="geometry", crs=gdf.crs)

    out_path = upload_dir / "final" / "quadras_segmentos_deflex.gpkg"
    gdf_out.to_file(out_path, driver="GPKG", encoding="utf-8")

    print(f"✅ Deflexões calculadas e salvas em: {out_path}")
    return gdf_out

def carregar_quadras_poligono(upload_dir: Path):
    """
    Carrega a camada de quadras como POLÍGONOS e garante que possui CRS.
    Tenta automaticamente os arquivos comuns do pipeline.
    """

    candidatos = [
        upload_dir / "quadras" / "quadras.shp",
        upload_dir / "quadras" / "quadras_m2s.shp",
        upload_dir / "quadras" / "quadras_dissolve.shp",
    ]

    for path in candidatos:
        if path.exists():
            print(f"🔎 Tentando carregar quadras: {path}")
            gdf = gpd.read_file(path)

            # Caso não tenha CRS
            if gdf.crs is None:
                print("⚠ Quadras sem CRS — aplicando EPSG:31983")
                gdf = gdf.set_crs(31983)

            # Geometria deve ser Polygon ou MultiPolygon
            if gdf.geom_type.isin(["Polygon", "MultiPolygon"]).any():
                print(f"Quadras carregadas com sucesso: {path}")
                return gdf

            else:
                print(f"⚠ Arquivo não contém polígonos (tipo: {gdf.geom_type.unique()})")

    raise FileNotFoundError(
        "Nenhuma camada válida de quadras poligonais foi encontrada."
    )

def gerar_memorial_quadras_docx(
        upload_dir: Path,
        arquivo_segmentos="quadras_segmentos.gpkg",
        nucleo="Teste",
        municipio="Teste",
        uf="Teste",
        promotor="Instituto Cidade Legal",
        saida_nome="memorial_quadras.docx"
    ):
    
    path_seg = upload_dir / "final" / arquivo_segmentos
    gdf = gpd.read_file(path_seg).sort_values(["quadra", "seq"])
    gdf_area = carregar_quadras_poligono(upload_dir)

    quadras = sorted(gdf["quadra"].unique(), key=lambda x: str(x))

    doc = Document()

    add_cabecalho_memorial_quadras(doc)

    # Estilo do documento
    estilo = doc.styles["Normal"]
    estilo.font.name = "Arial"
    estilo.font.size = Pt(11)

    for quadra in quadras:
        sub = gdf[gdf["quadra"] == quadra].copy().reset_index(drop=True)
        sub_poly = gdf_area[gdf_area["quadra"] == quadra].copy().reset_index(drop=True)
        # -----------------------------
        # CALCULAR ÁREA E PERÍMETRO DA QUADRA
        # -----------------------------
        geom_quadra = sub_poly.geometry.union_all()

        print("Quadra:", quadra)
        print("Tipo:", geom_quadra.geom_type)
        print("Is valid:", geom_quadra.is_valid)
        print("Área bruta:", geom_quadra.area)
        print("csr", gdf.crs)

        # Corrige geometrias quebradas
        if geom_quadra.is_empty:
            print(f"⚠ Quadra {quadra}: geometria vazia ao dissolver.")
        else:
            geom_quadra = geom_quadra.buffer(0)

        # Agora sim: área e perímetro verdadeiros
        area_quadra = geom_quadra.area
        print("AREA QUADRA:", area_quadra)
        perimetro_quadra = sub.union_all().length

        # -----------------------------
        # BLOCO: QUADRA / ÁREA / PERÍMETRO
        # -----------------------------
        add_bloco_info_quadra(
            doc,
            quadra=quadra,
            area_m2=area_quadra,
            perimetro_m=perimetro_quadra
        )

        doc.add_paragraph()

        # --------------------------
        # INÍCIO DO TEXTÃO
        # --------------------------

        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(20)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


        for i, row in sub.iterrows():

            P1 = f"P{str(i+1).zfill(2)}"
            P2 = f"P{str(i+2).zfill(2)}" if i < len(sub)-1 else "P01"

            x1, y1 = row["x1"], row["y1"]
            x2, y2 = row["x2"], row["y2"]
            dist = row["comprimento"]
            az = row["azimute"]
            conf = row["confronto"]

            coord1 = f"N {fmt_coord(y1)}m e E {fmt_coord(x1)}m"
            coord2 = f"N {fmt_coord(y2)}m e E {fmt_coord(x2)}m"

            az_dms = azimute_dms(az)
            dist_fmt = fmt_dist(dist)

            # Primeiro vértice
            if i == 0:
                texto = (
                    f"Inicia-se a descrição desta quadra no vértice {P1}, "
                    f"de coordenadas {coord1}; "
                    f"deste, segue confrontando com {conf}, "
                    f"com os seguintes azimutes e distâncias: "
                    f"{az_dms} e {dist_fmt} m até o vértice {P2}, "
                    f"de coordenadas {coord2}; "
                )
            elif i == len(sub)-1:
                texto = (
                    f" {az_dms} e {dist_fmt} m até o vértice P01, "
                    f"ponto inicial da descrição deste perímetro."
                )
            else:
                texto = (
                    f"{az_dms} e {dist_fmt} m até o vértice {P2}, "
                    f"de coordenadas {coord2}; "
                )

            run = p.add_run(texto)

        # --------------------------
        # Rodapé padrão da norma
        # --------------------------

        doc.add_paragraph()
        rod = doc.add_paragraph()
        rod.paragraph_format.first_line_indent = Pt(20)
        rod.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        rod.add_run(
            "Todas as coordenadas aqui descritas estão georreferenciadas ao Sistema Geodésico "
            "Brasileiro, de coordenadas N m e E m, e encontram-se representadas no Sistema "
            "U T M, referenciadas ao Meridiano Central n° 39º00', fuso -24, tendo como datum o "
            "SIRGAS2000. Todos os azimutes, distâncias, área e perímetro foram calculados no "
            "plano de projeção U T M."
        )

        doc.add_page_break()

    # --------------------------
    # Salvar DOCX final
    # --------------------------
    out_path = upload_dir / "memoriais" / saida_nome
    out_path.parent.mkdir(exist_ok=True, parents=True)
    doc.save(str(out_path))

    print("Memorial de quadras gerado com sucesso!")
    print("Arquivo:", out_path)
    return out_path